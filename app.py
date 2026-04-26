import streamlit as st
import requests
import re
import time
import zipfile
import shutil
import copy
import io
import calendar
import tempfile
import os
import subprocess
from datetime import date
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

st.set_page_config(
    page_title="Gerador de RPCM",
    page_icon="⚕️",
    layout="centered",
)

# ─── Identidade visual ──────────────────────────────────────────────────────
# Tema "Saúde Militar" — verde-oliva + dourado-arenoso, sem elementos oficiais
COR_PRIMARIA = "#3D4F2D"
COR_PRIMARIA_CLARA = "#4A5F38"
COR_DOURADO = "#B8985A"
COR_DOURADO_CLARO = "#D4B872"
COR_PAPEL = "#F8F5EE"
COR_PAPEL_ESCURO = "#EDE8DA"
COR_TEXTO = "#2C2C2C"
COR_TEXTO_FRACO = "#6B6B6B"
COR_ERRO = "#8B3A3A"

CSS_GLOBAL = f"""
<style>
.main .block-container {{
    padding-top: 1.4rem;
    padding-bottom: 3rem;
    max-width: 760px;
}}

/* Header banner */
.app-header {{
    background: linear-gradient(135deg, {COR_PRIMARIA} 0%, {COR_PRIMARIA_CLARA} 100%);
    color: {COR_PAPEL};
    padding: 1.6rem 1.8rem 1.3rem;
    border-radius: 10px;
    margin-bottom: 0;
    box-shadow: 0 4px 14px rgba(40, 50, 25, 0.18);
    display: flex;
    align-items: center;
    gap: 1.1rem;
}}
.app-header .logo {{ flex-shrink: 0; }}
.app-header h1 {{
    margin: 0;
    color: {COR_PAPEL};
    font-size: 1.55rem;
    font-weight: 600;
    letter-spacing: 0.4px;
    line-height: 1.2;
}}
.app-header p {{
    margin: 0.25rem 0 0;
    color: #D9CFB5;
    font-size: 0.88rem;
    font-weight: 400;
}}

.gold-divider {{
    height: 3px;
    background: linear-gradient(90deg, {COR_DOURADO} 0%, {COR_DOURADO_CLARO} 50%, {COR_DOURADO} 100%);
    margin: 0 0 1.6rem 0;
    border-radius: 2px;
}}

/* Cabeçalhos de seção numerados */
.section-header {{
    color: {COR_PRIMARIA};
    font-size: 1.05rem;
    font-weight: 600;
    margin: 1.6rem 0 0.7rem;
    padding-bottom: 0.45rem;
    border-bottom: 1px solid {COR_DOURADO};
    display: flex;
    align-items: center;
    gap: 0.55rem;
    letter-spacing: 0.3px;
}}
.section-number {{
    background: {COR_PRIMARIA};
    color: {COR_PAPEL};
    width: 24px;
    height: 24px;
    border-radius: 50%;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    font-size: 0.82rem;
    font-weight: 700;
    flex-shrink: 0;
}}

/* Botões */
.stButton > button {{
    border-radius: 6px;
    font-weight: 500;
    transition: all 0.15s ease;
}}
.stButton > button[kind="primary"] {{
    background-color: {COR_PRIMARIA};
    border-color: {COR_PRIMARIA};
    color: {COR_PAPEL};
    box-shadow: 0 2px 6px rgba(61, 79, 45, 0.25);
}}
.stButton > button[kind="primary"]:hover {{
    background-color: {COR_PRIMARIA_CLARA};
    border-color: {COR_PRIMARIA_CLARA};
    box-shadow: 0 3px 10px rgba(61, 79, 45, 0.35);
}}
.stButton > button:not([kind="primary"]) {{
    border-color: {COR_DOURADO};
    color: {COR_PRIMARIA};
}}
.stButton > button:not([kind="primary"]):hover {{
    border-color: {COR_PRIMARIA};
    color: {COR_PRIMARIA};
    background-color: {COR_PAPEL_ESCURO};
}}

/* Caixas de mensagem */
div[data-testid="stAlert"] {{ border-radius: 6px; }}
div[data-testid="stFileUploader"] section {{
    border: 1px dashed {COR_DOURADO};
    border-radius: 8px;
    background-color: rgba(184, 152, 90, 0.05);
}}
div[data-testid="stFileUploader"] section:hover {{
    border-color: {COR_PRIMARIA};
    background-color: rgba(61, 79, 45, 0.04);
}}

/* Inputs e selectboxes */
input, .stTextInput input {{
    border-radius: 6px !important;
}}

/* Download button (saída final) */
div[data-testid="stDownloadButton"] button {{
    background-color: {COR_DOURADO} !important;
    border-color: {COR_DOURADO} !important;
    color: {COR_PRIMARIA} !important;
    font-weight: 600 !important;
}}
div[data-testid="stDownloadButton"] button:hover {{
    background-color: {COR_DOURADO_CLARO} !important;
    border-color: {COR_DOURADO_CLARO} !important;
}}

/* Footer */
.app-footer {{
    margin-top: 2.5rem;
    padding-top: 1.1rem;
    border-top: 1px solid {COR_DOURADO};
    color: {COR_TEXTO_FRACO};
    font-size: 0.78rem;
    text-align: center;
    line-height: 1.6;
}}
.app-footer strong {{ color: {COR_PRIMARIA}; }}

/* Esconde menu/footer Streamlit pra ficar mais limpo */
#MainMenu {{ visibility: hidden; }}
footer {{ visibility: hidden; }}
</style>
"""

LOGO_SVG = (
    f'<svg class="logo" width="58" height="58" viewBox="0 0 60 60" xmlns="http://www.w3.org/2000/svg">'
    f'<path d="M30 4 L52 11 V32 Q52 49 30 56 Q8 49 8 32 V11 Z" fill="{COR_PAPEL}" stroke="{COR_DOURADO}" stroke-width="2"/>'
    f'<rect x="26" y="17" width="8" height="24" fill="{COR_PRIMARIA}" rx="1"/>'
    f'<rect x="18" y="25" width="24" height="8" fill="{COR_PRIMARIA}" rx="1"/>'
    f'<circle cx="30" cy="29" r="2.2" fill="{COR_DOURADO}"/>'
    f'</svg>'
)

st.markdown(CSS_GLOBAL, unsafe_allow_html=True)

st.markdown(
    f"""<div class="app-header">{LOGO_SVG}<div><h1>Gerador de RPCM</h1><p>Relatório de Prestação de Contas Mensal &middot; Contratos de Credenciamento</p></div></div><div class="gold-divider"></div>""",
    unsafe_allow_html=True,
)

def _section_header(numero, titulo):
    """Renderiza um cabeçalho de seção numerado e estilizado."""
    st.markdown(
        f'<div class="section-header">'
        f'<span class="section-number">{numero}</span>{titulo}'
        f'</div>',
        unsafe_allow_html=True,
    )

MESES = {
    1:  ("JAN", "JANEIRO"),
    2:  ("FEV", "FEVEREIRO"),
    3:  ("MAR", "MARCO"),
    4:  ("ABR", "ABRIL"),
    5:  ("MAI", "MAIO"),
    6:  ("JUN", "JUNHO"),
    7:  ("JUL", "JULHO"),
    8:  ("AGO", "AGOSTO"),
    9:  ("SET", "SETEMBRO"),
    10: ("OUT", "OUTUBRO"),
    11: ("NOV", "NOVEMBRO"),
    12: ("DEZ", "DEZEMBRO"),
}
MESES_LISTA = [
    "Janeiro","Fevereiro","Março","Abril","Maio","Junho",
    "Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"
]

# ODT namespaces
NS_TEXT  = 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
NS_TABLE = 'urn:oasis:names:tc:opendocument:xmlns:table:1.0'

# ─── Funções de CNPJ ────────────────────────────────────────────────────────

def _formatar_cnpj(digits):
    """Formata 14 dígitos no padrão XX.XXX.XXX/XXXX-XX."""
    d = re.sub(r'\D', '', digits)
    if len(d) == 14:
        return f'{d[0:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:14]}'
    return None

def limpar_cnpj(cnpj):
    return re.sub(r'[.\-/]', '', cnpj)

def consultar_empresa(cnpj_limpo):
    """Consulta razão social na BrasilAPI. Retorna dict com dados ou None."""
    try:
        r = requests.get(
            f'https://brasilapi.com.br/api/cnpj/v1/{cnpj_limpo}',
            timeout=15,
        )
        if r.status_code == 200:
            data = r.json()
            return {
                'razao_social': data.get('razao_social', '') or '',
                'nome_fantasia': data.get('nome_fantasia', '') or '',
                'situacao': data.get('descricao_situacao_cadastral', '') or '',
            }
    except Exception:
        pass
    return None

def validar_cnpj_dv(cnpj):
    """Valida os 2 dígitos verificadores do CNPJ."""
    d = re.sub(r'\D', '', cnpj)
    if len(d) != 14 or len(set(d)) == 1:
        return False
    pesos1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma = sum(int(d[i]) * pesos1[i] for i in range(12))
    dv1 = 0 if soma % 11 < 2 else 11 - (soma % 11)
    if dv1 != int(d[12]):
        return False
    pesos2 = [6] + pesos1
    soma = sum(int(d[i]) * pesos2[i] for i in range(13))
    dv2 = 0 if soma % 11 < 2 else 11 - (soma % 11)
    return dv2 == int(d[13])

def extrair_cnpjs_texto(texto):
    """Extrai todos os CNPJs válidos (com DV correto) do texto. Retorna lista
    de strings de 14 dígitos limpos, na ordem em que aparecem, sem duplicatas."""
    padrao = re.compile(r'\d{2}[\s.\-/]*\d{3}[\s.\-/]*\d{3}[\s.\-/]*\d{4}[\s.\-/]*\d{2}')
    encontrados = []
    vistos = set()
    for cand in padrao.findall(texto):
        limpo = re.sub(r'\D', '', cand)
        if len(limpo) == 14 and limpo not in vistos and validar_cnpj_dv(limpo):
            vistos.add(limpo)
            encontrados.append(limpo)
    return encontrados

def extrair_nome_ocs(texto):
    """Tenta extrair o nome da empresa a partir do campo 'OCS:' do documento.
    Retorna string ou None."""
    m = re.search(
        r'OCS\s*:?\s*([^.\n,]+?)(?:\.|\bCNPJ\b|\n|$)',
        texto,
        re.IGNORECASE,
    )
    if m:
        nome = m.group(1).strip()
        if nome:
            return nome
    return None

def extrair_texto_documento(file_bytes, filename):
    """Extrai todo o texto de um documento (.docx/.dotx/.odt/.doc).
    Para .doc, usa LibreOffice para converter antes."""
    ext = os.path.splitext(filename)[1].lower()

    if ext in ('.docx', '.dotx'):
        doc = abrir_documento(file_bytes, filename)
        partes = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    partes.append(cell.text)
        return '\n'.join(partes)

    if ext == '.odt':
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                content = z.read('content.xml')
            tree = etree.fromstring(content)
            partes = []
            for node in tree.iter():
                if node.text:
                    partes.append(node.text)
                if node.tail:
                    partes.append(node.tail)
            return ' '.join(partes)
        except Exception:
            return ''

    if ext == '.doc':
        try:
            tmp_dir = tempfile.mkdtemp()
            doc_path = os.path.join(tmp_dir, filename)
            with open(doc_path, 'wb') as f:
                f.write(file_bytes)
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx',
                 '--outdir', tmp_dir, doc_path],
                capture_output=True, timeout=90,
            )
            docx_path = os.path.join(tmp_dir, os.path.splitext(filename)[0] + '.docx')
            if os.path.exists(docx_path):
                with open(docx_path, 'rb') as f:
                    return extrair_texto_documento(f.read(), docx_path)
        except Exception:
            pass
        return ''

    return ''

# ─── API Portal da Transparência ────────────────────────────────────────────

def get_pagamentos(cnpj_limpo, mes_num, ano):
    """Busca pagamentos via API oficial do Portal da Transparência.

    A API só aceita filtro por `ano` (ano de **emissão do documento**, ou seja,
    da OB), não por mês/ano de pagamento. Como pagamentos de um determinado
    mês podem vir de OBs emitidas em anos anteriores (empenhos antigos),
    consultamos o ano de referência e os 2 anos anteriores. Em seguida,
    filtramos manualmente pela data efetiva do pagamento (campo `data` da
    resposta) para manter só os documentos pagos no mês de referência.

    Limite de paginação como salvaguarda contra loop em caso de bug da API.
    """
    api_key = st.secrets.get("TRANSPARENCIA_API_KEY", "")

    headers = {
        'chave-api-dados': api_key,
        'Accept': 'application/json',
    }

    # Cobre o ano informado e o ano anterior. Pagamentos de janeiro,
    # fevereiro e março costumam vir de empenhos emitidos no ano anterior.
    # Empenhos mais antigos (2+ anos) são raros — se aparecer caso assim,
    # o usuário poderá apontar e revemos a estratégia.
    anos_busca = [int(ano), int(ano) - 1]

    todos = []
    ultimo_status = None
    ultimo_erro = None
    MAX_PAGES = 50  # salvaguarda contra loop infinito

    URL = 'https://api.portaldatransparencia.gov.br/api-de-dados/despesas/documentos-por-favorecido'

    def _request_pagina(params):
        """Faz a chamada com 1 retry simples se a API responder 200 mas com
        body vazio/inválido — caso típico de rate limit ou carga temporária."""
        for tentativa in (1, 2):
            try:
                r = requests.get(URL, params=params, headers=headers, timeout=30)
                if r.status_code != 200:
                    return None, r.status_code, r.text[:300]
                try:
                    data = r.json()
                except ValueError:
                    if tentativa == 1:
                        time.sleep(1.5)
                        continue
                    return None, 200, f"JSON inválido (body: '{r.text[:120]}')"
                return data, 200, None
            except Exception as e:
                if tentativa == 1:
                    time.sleep(1.5)
                    continue
                return None, None, str(e)
        return None, None, 'falha após retries'

    for ano_busca in anos_busca:
        pagina = 1
        while pagina <= MAX_PAGES:
            params = {
                'codigoPessoa': cnpj_limpo,
                'fase': 3,  # 3 = Pagamento
                'ano': ano_busca,
                'pagina': pagina,
            }
            data, status, erro = _request_pagina(params)
            ultimo_status = status if status is not None else ultimo_status
            if erro is not None:
                ultimo_erro = erro
                break
            if not isinstance(data, list) or len(data) == 0:
                break
            todos.extend(data)
            if len(data) < 500:
                break
            pagina += 1
            time.sleep(0.25)  # respeita rate limit entre páginas
        time.sleep(0.4)  # respeita rate limit entre anos

    # Verificação adicional: confirma que cada item está no mês/ano pedido.
    # A API já filtra, mas mantemos a checagem como guarda de segurança.
    pagamentos = []
    mes_str = f'{mes_num:02d}'
    ano_str = str(ano)
    for item in todos:
        data_pgto = item.get('data', item.get('dataDocumento', ''))
        data_str  = str(data_pgto)
        mes_ok = False
        # BR: DD/MM/YYYY → mês em [3:5], ano em [6:10]
        if len(data_str) >= 10 and data_str[2:3] == '/' and data_str[3:5] == mes_str and data_str[6:10] == ano_str:
            mes_ok = True
        # ISO: YYYY-MM-DD → ano em [0:4], mês em [5:7]
        elif len(data_str) >= 10 and data_str[4:5] == '-' and data_str[0:4] == ano_str and data_str[5:7] == mes_str:
            mes_ok = True
        if not mes_ok:
            continue
        doc_num   = item.get('documentoResumido', '') or item.get('documento', '')
        valor_raw = item.get('valor', item.get('valorDocumento', '0'))
        try:
            if isinstance(valor_raw, str):
                v = float(valor_raw.replace('.', '').replace(',', '.'))
            else:
                v = float(valor_raw)
            pagamentos.append((doc_num, _normalizar_data_br(data_pgto), formatar_valor(v), v))
        except Exception:
            pass

    # Ordena por data (cronológica) e depois por documento — saída estável
    pagamentos.sort(key=lambda p: (_chave_data(p[1]), p[0]))

    return pagamentos, ultimo_status, ultimo_erro, todos

def _normalizar_data_br(data_str):
    """Normaliza data para DD/MM/YYYY (aceita ISO YYYY-MM-DD ou já BR)."""
    s = str(data_str).strip()
    if len(s) >= 10 and s[4:5] == '-':
        return f'{s[8:10]}/{s[5:7]}/{s[0:4]}'
    return s[:10] if len(s) >= 10 else s

def _chave_data(data_br):
    """Converte DD/MM/YYYY em chave ordenável (YYYYMMDD)."""
    s = str(data_br)
    if len(s) >= 10 and s[2:3] == '/' and s[5:6] == '/':
        return s[6:10] + s[3:5] + s[0:2]
    return s

def formatar_valor(v):
    """Formata float para padrão brasileiro: R$ 1.234,56"""
    return 'R$ {:,.2f}'.format(v).replace(',', 'X').replace('.', ',').replace('X', '.')

def calcular_total(pagamentos):
    total = sum(v for _, _, _, v in pagamentos)
    return '{:,.2f}'.format(total).replace(',', 'X').replace('.', ',').replace('X', '.')

# ─── DOCX / DOTX ────────────────────────────────────────────────────────────

def abrir_documento(file_bytes, filename):
    """Abre .docx ou .dotx com python-docx, convertendo .dotx se necessário."""
    ext = os.path.splitext(filename)[1].lower()
    tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
    tmp.write(file_bytes)
    tmp.close()
    try:
        return Document(tmp.name)
    except ValueError:
        # .dotx: corrigir content type
        dst = tmp.name.replace(ext, '.docx')
        shutil.copy2(tmp.name, dst)
        with zipfile.ZipFile(dst, 'r') as zin:
            files = {n: zin.read(n) for n in zin.namelist()}
        ct = files['[Content_Types].xml'].decode('utf-8')
        ct = ct.replace(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
        )
        files['[Content_Types].xml'] = ct.encode('utf-8')
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for n, d in files.items():
                zout.writestr(n, d)
        return Document(dst)

def atualizar_documento(doc, mes_abrev, ano, pagamentos, total_str):
    """Atualiza mês/ano no cabeçalho e substitui a tabela de pagamentos (DOCX)."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 1. Cabeçalho
    for para in doc.paragraphs:
        for run in para.runs:
            if re.search(r'\b(JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\b', run.text):
                run.text = re.sub(
                    r'\b(JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\b',
                    mes_abrev, run.text)
            if re.search(r'\b(JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)/\d{4}\b', run.text):
                run.text = re.sub(r'(?<=\/)\d{4}', ano, run.text)

    # 2. Tabela de pagamentos
    for table in doc.tables:
        if 'DOCUMENTO' not in table.rows[0].cells[0].text:
            continue

        tbl = table._tbl
        all_rows = list(table.rows)

        # Detecta linha de TOTAL: procura "TOTAL" em qualquer célula da linha
        # (mais robusto que olhar só a primeira — funciona com templates antigos
        # que têm "TOTAL PAGO" na col 1 e com o formato novo, em que a col 1
        # fica vazia e o label "Valor Total" vai na col 2)
        total_tr = None
        for row in all_rows[1:]:
            row_text = ' '.join(c.text for c in row.cells).upper()
            if 'TOTAL' in row_text:
                total_tr = row._tr
                break

        # Coleta as linhas de dados (tudo exceto cabeçalho e total)
        data_trs = [row._tr for row in all_rows[1:] if row._tr is not total_tr]

        # Usa a 1ª linha de dados como TEMPLATE de estilo; sem ela, cai no header
        if data_trs:
            template_tr = copy.deepcopy(data_trs[0])
        else:
            template_tr = copy.deepcopy(all_rows[0]._tr)

        # Remove todas as linhas de dados existentes (incluindo a de exemplo)
        for tr in data_trs:
            tbl.remove(tr)

        def set_tc_text(tc, texto):
            """Sobrescreve o texto de uma célula DOCX preservando o estilo do
            primeiro run (rPr) e a formatação do parágrafo."""
            para = tc.find(qn('w:p'))
            if para is None:
                para = etree.SubElement(tc, f'{{{W}}}p')
            old_r = para.find(qn('w:r'))
            rPr = old_r.find(qn('w:rPr')) if old_r is not None else None
            for r in para.findall(qn('w:r')):
                para.remove(r)
            new_r = etree.SubElement(para, f'{{{W}}}r')
            if rPr is not None:
                new_r.insert(0, copy.deepcopy(rPr))
            t = etree.SubElement(new_r, f'{{{W}}}t')
            t.text = texto
            if texto and (texto.startswith(' ') or texto.endswith(' ')):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        def make_row(doc_num, data, valor):
            new_tr = copy.deepcopy(template_tr)
            tcs = new_tr.findall(qn('w:tc'))
            for tc, texto in zip(tcs, [doc_num, data, valor]):
                set_tc_text(tc, texto)
            return new_tr

        def aplicar_negrito_celula(tc):
            for r in tc.iter(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.Element(f'{{{W}}}rPr')
                    r.insert(0, rPr)
                if rPr.find(qn('w:b')) is None:
                    etree.SubElement(rPr, f'{{{W}}}b')

        if total_tr is not None:
            for doc_num, data, valor, _ in pagamentos:
                total_tr.addprevious(make_row(doc_num, data, valor))

            # Reescreve a linha do total no formato Opção B:
            # col 1 vazia | col 2: "Valor Total" | col 3: "R$ X,XX"
            tcs = total_tr.findall(qn('w:tc'))
            if len(tcs) >= 3:
                # Já tem 3+ células — sobrescreve as 3 primeiras, mantém formatação
                set_tc_text(tcs[0], '')
                set_tc_text(tcs[1], 'Valor Total')
                set_tc_text(tcs[2], f'R$ {total_str}')
                for extra in tcs[3:]:
                    set_tc_text(extra, '')
            else:
                # Linha tem células mescladas (ex.: gridSpan=2). Substitui por
                # uma linha nova baseada na linha de dados (3 células) e aplica
                # negrito (típico de linha de total)
                new_total_tr = make_row('', 'Valor Total', f'R$ {total_str}')
                for tc in new_total_tr.findall(qn('w:tc')):
                    aplicar_negrito_celula(tc)
                total_tr.addprevious(new_total_tr)
                tbl.remove(total_tr)
        else:
            # Sem linha de TOTAL no template — anexa pagamentos e cria linha de total
            for doc_num, data, valor, _ in pagamentos:
                tbl.append(make_row(doc_num, data, valor))
            new_total_tr = make_row('', 'Valor Total', f'R$ {total_str}')
            for tc in new_total_tr.findall(qn('w:tc')):
                aplicar_negrito_celula(tc)
            tbl.append(new_total_tr)
        break

    return doc

# ─── ODT ────────────────────────────────────────────────────────────────────

def _odt_cell_text(cell):
    """Extrai texto de uma célula ODT."""
    partes = []
    for node in cell.iter():
        if node.text:
            partes.append(node.text)
        if node.tail:
            partes.append(node.tail)
    return ''.join(partes).strip()

def _odt_set_cell_text(cell, texto):
    """Define texto de uma célula ODT preservando o estilo de parágrafo
    (text:style-name no <text:p>) e o estilo de texto (<text:span>) que
    estiverem na célula. Sem isso, a linha inserida sai com formatação
    diferente do modelo."""
    NS_T = f'{{{NS_TEXT}}}'
    paragrafos = cell.findall(f'{NS_T}p')
    if paragrafos:
        p = paragrafos[0]
        # Tenta reaproveitar um <text:span> existente (mantém estilo de texto)
        spans = p.findall(f'{NS_T}span')
        if spans:
            span = spans[0]
            # Limpa filhos e tail do span, mantém atributos (estilo)
            for child in list(span):
                span.remove(child)
            span.text = texto
            span.tail = None
            # Remove qualquer conteúdo solto no parágrafo (texto direto, outros spans)
            p.text = None
            for child in list(p):
                if child is not span:
                    p.remove(child)
        else:
            # Sem span — define texto direto no parágrafo, mantém estilo do <text:p>
            for child in list(p):
                p.remove(child)
            p.text = texto
        # Remove parágrafos extras (se houver)
        for extra in paragrafos[1:]:
            cell.remove(extra)
    else:
        p = etree.SubElement(cell, f'{NS_T}p')
        p.text = texto

def _substituir_mes_ano_odt(tree, mes_abrev, ano):
    """Substitui mês/ano em todos os nós de texto da árvore XML do ODT."""
    meses_re = r'\b(JAN|FEV|MAR|ABR|MAI|JUN|JUL|AGO|SET|OUT|NOV|DEZ)\b'
    for node in tree.iter():
        if node.text and re.search(meses_re, node.text):
            node.text = re.sub(meses_re, mes_abrev, node.text)
            node.text = re.sub(r'(?<=/)\d{4}', ano, node.text)
        if node.tail and re.search(meses_re, node.tail):
            node.tail = re.sub(meses_re, mes_abrev, node.tail)
            node.tail = re.sub(r'(?<=/)\d{4}', ano, node.tail)

def atualizar_odt(file_bytes, mes_abrev, ano, pagamentos, total_str):
    """Atualiza ODT: substitui mês/ano e reconstrói tabela de pagamentos."""
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        files = {n: z.read(n) for n in z.namelist()}

    NS_T = f'{{{NS_TEXT}}}'
    NS_TB = f'{{{NS_TABLE}}}'

    # Processa content.xml
    content_xml = files.get('content.xml')
    if content_xml is None:
        raise ValueError("ODT inválido: content.xml não encontrado")

    tree = etree.fromstring(content_xml)

    # 1. Substituir mês/ano
    _substituir_mes_ano_odt(tree, mes_abrev, ano)

    # 2. Atualizar tabela de pagamentos
    for table in tree.iter(f'{NS_TB}table'):
        rows = table.findall(f'{NS_TB}table-row')
        if len(rows) < 2:
            continue
        header_cells = rows[0].findall(f'{NS_TB}table-cell')
        if not header_cells:
            continue
        if 'DOCUMENTO' not in _odt_cell_text(header_cells[0]).upper():
            continue

        # Detecta linha de TOTAL: procura "TOTAL" em qualquer célula da linha
        # (suporta tanto template antigo quanto o formato Opção B com col 1 vazia)
        total_row = None
        for r in rows[1:]:
            cells = r.findall(f'{NS_TB}table-cell')
            row_text = ' '.join(_odt_cell_text(c) for c in cells).upper()
            if 'TOTAL' in row_text:
                total_row = r
                break

        # Coleta linhas de dados (todas exceto cabeçalho e total)
        data_rows = [r for r in rows[1:] if r is not total_row]

        # Define um "template de linha" — preferimos a 1ª linha de dados existente
        # (preserva estilos de célula e <text:span> do modelo do usuário).
        if data_rows:
            template_row = copy.deepcopy(data_rows[0])
        else:
            # Sem linha de exemplo no template: usa o cabeçalho como base estrutural
            template_row = copy.deepcopy(rows[0])

        # Remove TODAS as linhas de dados existentes (incluindo a de exemplo)
        for r in data_rows:
            table.remove(r)

        def _make_row(doc_num, data, valor):
            new_row = copy.deepcopy(template_row)
            cells = new_row.findall(f'{NS_TB}table-cell')
            if len(cells) >= 3:
                _odt_set_cell_text(cells[0], str(doc_num))
                _odt_set_cell_text(cells[1], str(data))
                _odt_set_cell_text(cells[2], str(valor))
            return new_row

        # Insere as novas linhas de pagamento e reescreve a linha do total
        if total_row is not None:
            for doc_num, data, valor, _ in pagamentos:
                total_row.addprevious(_make_row(doc_num, data, valor))

            # Reescreve a linha do total no formato Opção B:
            # col 1 vazia | col 2: "Valor Total" | col 3: "R$ X,XX"
            total_cells = total_row.findall(f'{NS_TB}table-cell')
            if len(total_cells) >= 3:
                _odt_set_cell_text(total_cells[0], '')
                _odt_set_cell_text(total_cells[1], 'Valor Total')
                _odt_set_cell_text(total_cells[2], f'R$ {total_str}')
                for extra in total_cells[3:]:
                    _odt_set_cell_text(extra, '')
            else:
                # Linha tem células mescladas — substitui por uma nova baseada
                # na linha de dados (3 células com mesma formatação)
                new_total_row = _make_row('', 'Valor Total', f'R$ {total_str}')
                total_row.addprevious(new_total_row)
                table.remove(total_row)
        else:
            # Sem linha de TOTAL no template — anexa pagamentos e cria linha de total
            for doc_num, data, valor, _ in pagamentos:
                table.append(_make_row(doc_num, data, valor))
            table.append(_make_row('', 'Valor Total', f'R$ {total_str}'))
        break

    files['content.xml'] = etree.tostring(
        tree, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # Atualiza styles.xml se houver (para cabeçalhos/rodapés)
    if 'styles.xml' in files:
        styles_tree = etree.fromstring(files['styles.xml'])
        _substituir_mes_ano_odt(styles_tree, mes_abrev, ano)
        files['styles.xml'] = etree.tostring(
            styles_tree, xml_declaration=True, encoding='UTF-8', standalone=True
        )

    # Remonta o ZIP (mimetype deve ser primeiro e sem compressão)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        if 'mimetype' in files:
            info = zipfile.ZipInfo('mimetype')
            info.compress_type = zipfile.ZIP_STORED
            zout.writestr(info, files['mimetype'])
        for n, d in files.items():
            if n == 'mimetype':
                continue
            zout.writestr(n, d)
    buf.seek(0)
    return buf.read()

# ─── DOC (legado) ───────────────────────────────────────────────────────────

def processar_doc_libreoffice(file_bytes, filename, mes_abrev, ano, pagamentos, total_str):
    """Tenta processar .doc via LibreOffice: converte para docx, processa, converte de volta."""
    tmp_dir = tempfile.mkdtemp()
    doc_path = os.path.join(tmp_dir, filename)
    with open(doc_path, 'wb') as f:
        f.write(file_bytes)

    # .doc → .docx
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', tmp_dir, doc_path],
        capture_output=True, timeout=90
    )
    docx_path = os.path.join(tmp_dir, os.path.splitext(filename)[0] + '.docx')
    if not os.path.exists(docx_path):
        raise RuntimeError("LibreOffice não converteu o arquivo")

    with open(docx_path, 'rb') as f:
        docx_bytes = f.read()

    doc = abrir_documento(docx_bytes, docx_path)
    doc = atualizar_documento(doc, mes_abrev, ano, pagamentos, total_str)

    processed_path = os.path.join(tmp_dir, 'saida.docx')
    doc.save(processed_path)

    # .docx → .doc
    subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'doc', '--outdir', tmp_dir, processed_path],
        capture_output=True, timeout=90
    )
    out_doc = os.path.join(tmp_dir, 'saida.doc')
    if os.path.exists(out_doc):
        with open(out_doc, 'rb') as f:
            return f.read(), '.doc'

    # Fallback: retorna como .docx
    with open(processed_path, 'rb') as f:
        return f.read(), '.docx'

# ─── Utilitário ─────────────────────────────────────────────────────────────

def nome_saida(nome_entrada, mes_nome_arquivo, ext_saida=None):
    """Deriva nome do arquivo de saída substituindo o mês."""
    base = os.path.splitext(nome_entrada)[0]
    ext  = ext_saida if ext_saida else os.path.splitext(nome_entrada)[1]
    meses_re = '|'.join([
        'JANEIRO','FEVEREIRO','MARCO','MARÇO','ABRIL','MAIO','JUNHO',
        'JULHO','AGOSTO','SETEMBRO','OUTUBRO','NOVEMBRO','DEZEMBRO',
        'JAN','FEV','MAR','ABR','MAI','JUN','JUL','AGO','SET','OUT','NOV','DEZ'
    ])
    novo = re.sub(
        rf'\b({meses_re})\b', mes_nome_arquivo,
        base.upper(), flags=re.IGNORECASE
    )
    return novo + ext

# ─── Interface ──────────────────────────────────────────────────────────────

# Estado da sessão
for _k, _v in [
    ('cnpj_confirmado', None),
    ('empresa_info', None),
    ('arquivo_chave', None),
    ('file_bytes', None),
    ('cnpj_extraido', None),
    ('nome_empresa_ocs', None),
    ('modo_manual', False),
]:
    if _k not in st.session_state:
        st.session_state[_k] = _v

_section_header(1, "Documento base")
st.caption(
    "Envie o RPCM modelo. O CNPJ é extraído automaticamente; "
    "se não for possível, você poderá digitá-lo na próxima seção."
)

uploaded = st.file_uploader(
    "Selecione o arquivo (.docx, .dotx, .odt ou .doc)",
    type=["docx", "dotx", "odt", "doc"],
    label_visibility="collapsed",
)

# Detecta upload novo e roda extração automática
if uploaded is not None:
    _bytes_atual = uploaded.getvalue()
    _chave_atual = f"{uploaded.name}|{len(_bytes_atual)}"

    if st.session_state.arquivo_chave != _chave_atual:
        # Arquivo novo — reseta tudo e tenta extrair CNPJ
        st.session_state.arquivo_chave = _chave_atual
        st.session_state.file_bytes = _bytes_atual
        st.session_state.cnpj_confirmado = None
        st.session_state.empresa_info = None
        st.session_state.cnpj_extraido = None
        st.session_state.nome_empresa_ocs = None
        st.session_state.modo_manual = False

        with st.spinner("Lendo documento e extraindo CNPJ..."):
            try:
                texto_doc = extrair_texto_documento(_bytes_atual, uploaded.name)
                st.session_state.nome_empresa_ocs = extrair_nome_ocs(texto_doc)
                cnpjs = extrair_cnpjs_texto(texto_doc)
                if cnpjs:
                    # CNPJ encontrado e validado pelo DV — auto-confirma sem
                    # exigir clique do usuário. Se quiser trocar, usa o botão.
                    st.session_state.cnpj_extraido = cnpjs[0]
                    st.session_state.cnpj_confirmado = _formatar_cnpj(cnpjs[0])
                    info_auto = consultar_empresa(cnpjs[0])
                    if info_auto:
                        st.session_state.empresa_info = info_auto
                else:
                    st.session_state.modo_manual = True
            except Exception as e:
                st.session_state.modo_manual = True
                st.error(f"Erro ao ler o documento: {e}")

# Bloco de CNPJ — depende do estado atual
if uploaded is not None:
    _section_header(2, "Empresa contratada")
    if st.session_state.cnpj_confirmado:
        # CNPJ confirmado (automático ou manual). Mostra dados e permite trocar.
        info = st.session_state.empresa_info or {}
        razao = info.get('razao_social', '')
        if st.session_state.cnpj_extraido:
            st.success(
                f"CNPJ identificado automaticamente no documento — "
                f"**{st.session_state.cnpj_confirmado}**"
            )
        else:
            st.success(
                f"CNPJ confirmado — **{st.session_state.cnpj_confirmado}**"
            )
        if info:
            nome_fantasia = info.get('nome_fantasia', '').strip()
            extra = ''
            if nome_fantasia and nome_fantasia.lower() != razao.strip().lower():
                extra = f"  \n**Nome fantasia:** {nome_fantasia}"
            st.markdown(
                f"**Empresa:** {razao}{extra}  \n"
                f"**Situação cadastral:** {info.get('situacao', '')}"
            )
        elif st.session_state.cnpj_extraido:
            st.caption(
                "Não foi possível consultar a BrasilAPI agora — os dados "
                "da empresa não estão disponíveis, mas o CNPJ foi validado."
            )
        if st.button("Trocar CNPJ", help="Usar outro CNPJ neste documento"):
            st.session_state.cnpj_confirmado = None
            st.session_state.empresa_info = None
            st.session_state.cnpj_extraido = None
            st.session_state.modo_manual = True
            st.rerun()

    else:
        # Modo manual: extração falhou ou usuário pediu pra trocar
        if not st.session_state.cnpj_extraido:
            msg = "Não foi possível extrair o CNPJ automaticamente do documento."
            if st.session_state.nome_empresa_ocs:
                msg += f"\n\n**Empresa identificada no documento:** {st.session_state.nome_empresa_ocs}"
            msg += "\n\nDigite o CNPJ manualmente abaixo:"
            st.warning(msg)

        cnpj_input = st.text_input(
            "CNPJ da empresa",
            placeholder="22.416.260/0001-85 ou apenas os 14 dígitos",
            max_chars=18,
        )
        cnpj_formatado_atual = _formatar_cnpj(cnpj_input) if cnpj_input.strip() else None

        if st.button("Buscar empresa", disabled=(cnpj_formatado_atual is None)):
            with st.spinner("Consultando BrasilAPI..."):
                info_m = consultar_empresa(limpar_cnpj(cnpj_formatado_atual))
            if info_m:
                st.session_state.empresa_info = info_m
            else:
                st.session_state.empresa_info = None
                st.error(
                    "Não foi possível consultar essa empresa. "
                    "Verifique o CNPJ ou tente novamente em alguns segundos."
                )

        if st.session_state.empresa_info and cnpj_formatado_atual:
            info = st.session_state.empresa_info
            nome_fantasia = info.get('nome_fantasia', '').strip()
            extra = ''
            if nome_fantasia and nome_fantasia.lower() != info['razao_social'].strip().lower():
                extra = f"  \n**Nome fantasia:** {nome_fantasia}"
            st.markdown(
                f"**Empresa encontrada:** {info['razao_social']}{extra}  \n"
                f"**Situação cadastral:** {info['situacao']}"
            )
            if st.button("Confirmar e usar este CNPJ", type="primary"):
                st.session_state.cnpj_confirmado = cnpj_formatado_atual
                st.rerun()

_section_header(3, "Período de referência")
col1, col2 = st.columns(2)
with col1:
    mes_selecionado = st.selectbox("Mês", MESES_LISTA)
with col2:
    ano_atual = date.today().year
    anos_disponiveis = list(range(ano_atual, ano_atual - 7, -1))
    ano_selecionado = st.selectbox("Ano", anos_disponiveis)
    ano_input = str(ano_selecionado)

# Cálculo do mês anterior à data atual (rola pra dezembro/ano-1 em janeiro)
hoje = date.today()
if hoje.month == 1:
    mes_anterior_num, ano_anterior_num = 12, hoje.year - 1
else:
    mes_anterior_num, ano_anterior_num = hoje.month - 1, hoje.year
label_mes_anterior = MESES_LISTA[mes_anterior_num - 1]

_section_header(4, "Geração do relatório")
botoes_disabled = (uploaded is None) or (st.session_state.cnpj_confirmado is None)
if botoes_disabled:
    st.caption(
        "_Aguardando o documento e a confirmação do CNPJ para liberar a geração._"
    )

botao_col1, botao_col2 = st.columns(2)
with botao_col1:
    gerar = st.button(
        "Gerar relatório",
        type="primary",
        disabled=botoes_disabled,
        use_container_width=True,
    )
with botao_col2:
    gerar_mes_anterior = st.button(
        f"Atalho — {label_mes_anterior}/{ano_anterior_num}",
        disabled=botoes_disabled,
        use_container_width=True,
        help="Gera o relatório referente ao mês anterior à data de hoje.",
    )

# ─── Lógica principal ───────────────────────────────────────────────────────

if (gerar or gerar_mes_anterior) and uploaded and st.session_state.cnpj_confirmado:
    if gerar_mes_anterior:
        mes_num = mes_anterior_num
        ano = str(ano_anterior_num)
        mes_selecionado = label_mes_anterior
    else:
        mes_num = MESES_LISTA.index(mes_selecionado) + 1
        ano = ano_input.strip()
    mes_abrev, mes_nome_arq = MESES[mes_num]
    ext_entrada = os.path.splitext(uploaded.name)[1].lower()

    progress = st.progress(0)
    status   = st.empty()
    file_bytes = st.session_state.file_bytes or uploaded.getvalue()

    # PASSO 1 — CNPJ (já confirmado pelo usuário)
    cnpj = st.session_state.cnpj_confirmado
    progress.progress(15)

    # PASSO 2 — Pagamentos via API oficial
    status.info(f"Buscando pagamentos de {mes_selecionado}/{ano} no Portal da Transparência...")
    pagamentos, api_status, api_erro, todos_brutos = get_pagamentos(limpar_cnpj(cnpj), mes_num, ano)
    progress.progress(65)

    if len(pagamentos) == 0:
        msg = f"Nenhum pagamento encontrado para {mes_selecionado}/{ano}. O relatório será gerado com tabela vazia."
        if api_status and api_status != 200:
            msg += f"\n\n_API retornou status **{api_status}**._"
            if api_erro:
                msg += f" Resposta: `{api_erro}`"
        elif api_status is None:
            msg += "\n\n_Não foi possível conectar à API (timeout ou erro de rede)._"
        st.warning(msg)
        total_str = "0,00"
    else:
        total_str = calcular_total(pagamentos)
        st.info(f"**{len(pagamentos)} pagamento(s)** localizados · Total: **R$ {total_str}**")

    # PASSO 3 — Gerar documento
    status.info("Atualizando documento...")
    ext_saida = ext_entrada
    mime_type = 'application/octet-stream'

    if ext_entrada == '.odt':
        output_bytes = atualizar_odt(file_bytes, mes_abrev, ano, pagamentos, total_str)
        ext_saida = '.odt'
        mime_type = 'application/vnd.oasis.opendocument.text'

    elif ext_entrada == '.doc':
        # Tenta LibreOffice; se não disponível, orienta o usuário
        try:
            output_bytes, ext_saida = processar_doc_libreoffice(
                file_bytes, uploaded.name, mes_abrev, ano, pagamentos, total_str
            )
            if ext_saida == '.docx':
                st.warning(
                    "Não foi possível manter o formato `.doc`. "
                    "O arquivo foi salvo como `.docx`."
                )
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        except Exception:
            st.error(
                "**Formato `.doc` não suportado para processamento automático.**\n\n"
                "O formato `.doc` (Word 97-2003) requer conversão prévia. "
                "Por favor, abra o arquivo no Word e salve como **`.docx`**, "
                "depois envie novamente."
            )
            st.stop()

    else:
        # DOCX / DOTX
        doc = abrir_documento(file_bytes, uploaded.name)
        doc = atualizar_documento(doc, mes_abrev, ano, pagamentos, total_str)
        buf = io.BytesIO()
        doc.save(buf)
        output_bytes = buf.getvalue()
        ext_saida = '.docx'
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    progress.progress(90)
    output_name = nome_saida(uploaded.name, mes_nome_arq, ext_saida)
    progress.progress(100)
    status.success("Documento gerado.")

    st.download_button(
        label=f"⬇  Baixar {output_name}",
        data=output_bytes,
        file_name=output_name,
        mime=mime_type,
        type="primary",
        use_container_width=True,
    )

    if pagamentos:
        st.markdown(
            f'<div class="section-header" style="margin-top:1.6rem">'
            f'<span class="section-number">5</span>Pagamentos incluídos'
            f'</div>',
            unsafe_allow_html=True,
        )
        import pandas as pd
        rows = [(d, dt, v) for d, dt, v, _ in pagamentos]
        rows.append(("", "Valor Total", f"R$ {total_str}"))
        df = pd.DataFrame(rows, columns=["Documento", "Data", "Valor"])
        st.dataframe(df, hide_index=True, use_container_width=True)

    # Detalhes técnicos (debug) — discreto, ao final
    with st.expander("Detalhes técnicos da requisição (avançado)"):
        st.write(f"**Status HTTP:** {api_status}")
        st.write(f"**Total de registros retornados pela API:** {len(todos_brutos)}")
        if api_erro:
            st.write(f"**Erro:** {api_erro}")
        if todos_brutos:
            st.write("**Primeiros registros (raw):**")
            st.json(todos_brutos[:5])

# ─── Rodapé ─────────────────────────────────────────────────────────────────

st.markdown(
    '<div class="app-footer"><strong>Aplicativo não oficial.</strong> Uso interno.<br>'
    'Dados públicos do Portal da Transparência do Governo Federal &middot; UG 167399</div>',
    unsafe_allow_html=True,
)
